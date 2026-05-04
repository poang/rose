import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import datetime

def add_table_with_style(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

def generate_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_heading("DKSDK V2 接口集成手册 (更新版)", 0)
    doc.add_paragraph(f"生成日期: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    
    # --- 第6章 典型应用流程 (完全重写为对外接口视角) ---
    doc.add_heading("第6章  典型应用流程", level=1)
    doc.add_paragraph("本章描述 SDK 的典型集成场景。流程聚焦于 MCU 与 SDK 公开接口的交互，屏蔽内部实现细节。")

    # 6.1 初始化
    doc.add_heading("6.1  系统初始化流程", level=2)
    add_table_with_style(doc, ["步骤", "操作", "说明"], [
        ["1", "注册回调", "调用 Sdk_Register* 系列函数注册 APDU、发送、结果通知等回调"],
        ["2", "调用 Sdk_Init()", "SDK 初始化内部资源并自动选择 SE Applet"],
        ["3", "启动调度", "MCU 必须以 1ms 周期调用 Sdk_routine() 驱动 SDK 运行"],
        ["4", "获取版本", "SDK 异步触发 get_version_cb 通知版本号和 SEID"]
    ], col_widths=[1.5, 4.0, 9.5])

    # 6.2 认证
    doc.add_heading("6.2  数字钥匙认证流程", level=2)
    doc.add_paragraph("[此处建议插入认证流程图]")
    add_table_with_style(doc, ["阶段", "动作", "MCU 职责"], [
        ["发起", "Sdk_Auth", "指定通道并开始认证"],
        ["SE 交互", "apdu_cb / Sdk_ApduCallback", "将 SDK 指令发给 SE，并将响应回传给 SDK"],
        ["无线交互", "send_cb / Sdk_SendCallback", "将 SDK 指令发给手机，并将响应回传给 SDK"],
        ["完成", "notify_auth_cb", "接收最终认证结果（成功含 CardID，失败含错误码）"]
    ], col_widths=[2.5, 4.5, 8.0])

    # 6.3 数据上报
    doc.add_heading("6.3  数据上报流程", level=2)
    add_table_with_style(doc, ["步骤", "动作", "说明"], [
        ["1", "Sdk_DataReport", "MCU 发起上报（16字节数据）"],
        ["2", "加密配合", "通过 apdu_cb/Sdk_ApduCallback 配合 SDK 完成数据加密"],
        ["3", "执行发送", "通过 send_cb 将加密数据发往手机"],
        ["4", "反馈通知", "通过 data_report_cb 接收手机端的确认结果"]
    ], col_widths=[1.5, 4.0, 9.5])

    # 6.4 RKE 流程
    doc.add_heading("6.4  RKE 遥控钥匙流程", level=2)
    doc.add_paragraph("[此处建议插入 RKE 流程图]")
    add_table_with_style(doc, ["步骤", "动作", "说明"], [
        ["1", "接收密文", "MCU 收到手机 RKE 数据，调用 Sdk_SendCallback 传给 SDK"],
        ["2", "解密配合", "配合 SDK 完成 SE 解密操作"],
        ["3", "业务处理", "rke_cb 触发，MCU 执行车辆动作（如开锁）"],
        ["4", "结果反馈", "MCU 调用 Sdk_RKECallback 告知结果，SDK 自动加密并发送回手机"]
    ], col_widths=[1.5, 4.0, 9.5])

    # 6.10 注意事项
    doc.add_heading("6.10 重要注意事项", level=2)
    doc.add_paragraph("1. 回调函数必须为非阻塞模式，严禁在回调内进行耗时操作。", style='List Bullet')
    doc.add_paragraph("2. Sdk_routine() 必须保持 1ms 的稳定调用频率。", style='List Bullet')
    doc.add_paragraph("3. 回调内的数据缓冲区在函数返回后可能失效，必须手动拷贝。", style='List Bullet')
    doc.add_paragraph("4. 同一通道上业务互斥（如认证时不可上报数据）。", style='List Bullet')
    doc.add_paragraph("5. 物理连接断开时，请务必调用 Sdk_Release 释放资源。", style='List Bullet')

    output_path = os.path.join(os.path.dirname(__file__), "DKSDK接口文档_20260504_Final.docx")
    doc.save(output_path)
    print(f"DONE: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_document()
