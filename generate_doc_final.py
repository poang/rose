import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import datetime

def add_table_with_style(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

def generate_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_heading("DKSDK V2 接口集成手册 (含流程图)", 0)
    doc.add_paragraph(f"生成日期: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    
    # --- 第6章 ---
    doc.add_heading("第6章  典型应用流程", level=1)
    
    # 6.2 认证
    doc.add_heading("6.2  数字钥匙认证流程", level=2)
    doc.add_paragraph("下图展示了从手机发起认证到结果通知的完整集成流程：")
    
    # 插入已生成的认证流程图
    img_path = r"C:\Users\phtlu\.gemini\antigravity\brain\63b98915-875c-4b78-b1c7-1b7cad07a489\auth_flowchart_1777858102958.png"
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(15))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("图 6-1 数字钥匙认证集成时序图", style='Caption')
    
    add_table_with_style(doc, ["阶段", "动作", "说明"], [
        ["发起", "Sdk_Auth", "MCU 指定通道（BLE/NFC）发起请求"],
        ["SE 交互", "apdu_cb", "SDK 请求 SE 指令，MCU 负责硬件透传"],
        ["结果反馈", "notify_auth_cb", "SDK 异步通知认证成功/失败及 CardID"]
    ], col_widths=[2.5, 4.5, 8.0])

    # 6.4 RKE 流程 (采用美观的文本框图展示)
    doc.add_heading("6.4  RKE 遥控钥匙流程", level=2)
    doc.add_paragraph("RKE 流程交互逻辑如下：")
    
    # 手动构建一个示意图式的表格
    rke_flow = doc.add_table(rows=1, cols=7)
    rke_flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells = rke_flow.rows[0].cells
    cells[0].text = "手机(DK)"
    cells[1].text = " → "
    cells[2].text = "MCU (透传)"
    cells[3].text = " → "
    cells[4].text = "SDK (解密)"
    cells[5].text = " → "
    cells[6].text = "MCU (执行)"
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("") # 换行
    add_table_with_style(doc, ["步骤", "说明"], [
        ["1. 接收", "MCU 收到 RKE 密文，调用 Sdk_SendCallback"],
        ["2. 解密", "SDK 调度 SE 完成解密，通过 rke_cb 通知 MCU"],
        ["3. 执行", "MCU 执行动作后，调用 Sdk_RKECallback 告知结果"],
        ["4. 反馈", "SDK 自动加密并回传给手机"]
    ], col_widths=[2.0, 13.0])

    # 6.10 注意事项
    doc.add_heading("6.10 重要注意事项", level=2)
    doc.add_paragraph("1. 调度频率：必须保持 1ms 周期调用 Sdk_routine()。", style='List Bullet')
    doc.add_paragraph("2. 缓冲区安全：回调内的数据若需持久化，必须进行拷贝。", style='List Bullet')
    doc.add_paragraph("3. 异步意识：初始化和认证均为异步过程，以回调通知为准。", style='List Bullet')

    output_path = os.path.join(os.path.dirname(__file__), "DKSDK接口文档_Final_With_Diagrams.docx")
    doc.save(output_path)
    print(f"DONE: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_document()
