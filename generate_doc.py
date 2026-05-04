"""
DKSDK 文档生成脚本
根据 DKSdk.h 和 DKSdk.c 源码，生成 Word 格式的 SDK 文档
包含带目录的完整文档结构
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ========================================================================
# 辅助函数
# ========================================================================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code_text, font_size=8):
    """添加代码块"""
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    # Add shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "2F5496")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, "D6E4F0")
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    doc.add_paragraph()  # spacing after table
    return table


def add_function_detail(doc, name, prototype, brief, params, return_desc, notes=None):
    """添加函数详细说明（表格形式）"""
    doc.add_heading(name, level=3)
    
    # 构建详情表格的行数据
    detail_rows = [
        ("函数原型", prototype),
        ("功能描述", brief),
    ]
    
    if params:
        # 构建参数子表
        param_lines = []
        for pname, ptype, pdesc in params:
            param_lines.append(f"{pname}（{ptype}）：{pdesc}")
        detail_rows.append(("参数说明", "\n".join(param_lines)))
    
    detail_rows.append(("返回值", return_desc))
    
    if notes:
        detail_rows.append(("注意事项", notes))
    
    # 创建2列表格（字段名 | 内容）
    table = doc.add_table(rows=len(detail_rows), cols=2)
    table.style = 'Table Grid'
    
    for r, (field_name, content) in enumerate(detail_rows):
        # 字段名列
        cell0 = table.rows[r].cells[0]
        cell0.text = ""
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(field_name)
        run0.font.bold = True
        run0.font.size = Pt(9)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell0, "2F5496")
        run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        # 内容列
        cell1 = table.rows[r].cells[1]
        cell1.text = ""
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(content)
        run1.font.size = Pt(9)
        if r % 2 == 0:
            set_cell_shading(cell1, "F2F6FC")
    
    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(12.5)
    
    doc.add_paragraph()  # spacing after table


def add_callback_detail_table(doc, name, proto, desc, params, return_desc, notes=None):
    """添加MCU回调函数详细说明（表格形式，用于第5章）"""
    doc.add_heading(name, level=3)
    
    detail_rows = [
        ("类型定义", proto),
        ("功能描述", desc),
    ]
    
    if params:
        param_lines = []
        for pname, ptype, pdesc in params:
            param_lines.append(f"{pname}（{ptype}）：{pdesc}")
        detail_rows.append(("参数说明", "\n".join(param_lines)))
    
    detail_rows.append(("返回值", return_desc))
    
    if notes:
        detail_rows.append(("实现注意事项", notes))
    
    table = doc.add_table(rows=len(detail_rows), cols=2)
    table.style = 'Table Grid'
    
    for r, (field_name, content) in enumerate(detail_rows):
        cell0 = table.rows[r].cells[0]
        cell0.text = ""
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(field_name)
        run0.font.bold = True
        run0.font.size = Pt(9)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell0, "2F5496")
        run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        cell1 = table.rows[r].cells[1]
        cell1.text = ""
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(content)
        run1.font.size = Pt(9)
        if r % 2 == 0:
            set_cell_shading(cell1, "F2F6FC")
    
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(12.5)
    
    doc.add_paragraph()


def add_interface_summary_table(doc, headers, rows, col_widths):
    """添加接口汇总表格（用于第4章和第5章的接口概览）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
        set_cell_shading(cell, "2F5496")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
            if r % 2 == 0:
                set_cell_shading(cell, "D6E4F0")
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    doc.add_paragraph()  # spacing after table
    return table


# ========================================================================
# 主文档生成函数
# ========================================================================

def generate_document():
    doc = Document()
    
    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # ---- 设置样式 ----
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置标题样式
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '黑体'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if i == 1:
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        elif i == 3:
            heading_style.font.size = Pt(12)
            heading_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        elif i == 4:
            heading_style.font.size = Pt(11)
            heading_style.font.color.rgb = RGBColor(0x37, 0x5F, 0x9F)
    
    # ====================================================================
    # 封面
    # ====================================================================
    for _ in range(6):
        doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("DKSDK 软件开发工具包")
    run.font.name = '黑体'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    
    doc.add_paragraph()
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("接口文档")
    run.font.name = '黑体'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    for _ in range(4):
        doc.add_paragraph()
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.add_run("版本：V1.0").font.size = Pt(12)
    doc.add_paragraph()
    info_p2 = doc.add_paragraph()
    info_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p2.add_run("日期：2026年5月").font.size = Pt(12)
    
    doc.add_page_break()
    
    # ====================================================================
    # 目录页（Word会自动根据标题生成目录，此处留占位）
    # ====================================================================
    doc.add_heading("目  录", level=1)
    p = doc.add_paragraph()
    p.add_run("（请在Word中右键此处，选择\"更新域\"以生成自动目录，或使用\"引用 -> 目录\"插入自动目录）").font.size = Pt(9)
    doc.add_page_break()
    
    # ====================================================================
    # 第1章 概述
    # ====================================================================
    doc.add_heading("第1章  概述", level=1)
    
    doc.add_heading("1.1  编写目的", level=2)
    doc.add_paragraph(
        "本文档旨在详细描述 DKSDK（Digital Key Software Development Kit）软件开发工具包的接口定义、"
        "数据结构和使用方法，为 MCU 端开发者提供完整的集成参考。通过本文档，开发者可以了解："
    )
    doc.add_paragraph("（1）SDK 的整体架构与设计思想；", style='List Number')
    doc.add_paragraph("（2）SDK 与 MCU 之间的接口划分与调用关系；", style='List Number')
    doc.add_paragraph("（3）各结构体、枚举类型的定义与使用说明；", style='List Number')
    doc.add_paragraph("（4）各个 API 函数的功能、参数、返回值及注意事项。", style='List Number')
    
    doc.add_heading("1.2  适用范围", level=2)
    doc.add_paragraph(
        "本文档适用于基于 DKSDK 进行数字钥匙系统开发的 MCU 固件工程师。DKSDK 运行在 SE（Secure Element）"
        "安全芯片上，通过 APDU 指令与 DK（Digital Key）设备（如手机、卡片、手表等）进行交互，"
        "完成数字钥匙的认证、数据上报、远程钥匙执行（RKE）等功能。"
    )
    
    doc.add_heading("1.3  术语与缩写", level=2)
    add_table_with_style(doc, 
        ["术语/缩写", "全称", "说明"],
        [
            ["SDK", "Software Development Kit", "软件开发工具包，运行于SE端"],
            ["MCU", "Microcontroller Unit", "微控制器，SDK的宿主系统"],
            ["SE", "Secure Element", "安全芯片，SDK的运行载体"],
            ["DK", "Digital Key", "数字钥匙设备（手机、卡片、手表等）"],
            ["APDU", "Application Protocol Data Unit", "应用协议数据单元，SE与DK间的指令格式"],
            ["RKE", "Remote Keyless Entry", "远程无钥匙进入"],
            ["BLE", "Bluetooth Low Energy", "低功耗蓝牙通信通道"],
            ["NFC", "Near Field Communication", "近场通信通道"],
            ["GPD", "Get Processing Data", "获取处理数据"],
            ["TSM", "Trusted Service Manager", "可信服务管理"],
            ["TLV", "Tag-Length-Value", "标签-长度-值编码格式"],
            ["ICCE", "ICC Environment", "IC卡环境，此处指蓝牙帧协议"],
            ["CRC", "Cyclic Redundancy Check", "循环冗余校验"],
            ["SW", "Status Word", "状态字，APDU指令的响应状态码"],
        ],
        col_widths=[2.5, 4.5, 8.0]
    )
    
    # ====================================================================
    # 第2章 系统架构
    # ====================================================================
    doc.add_heading("第2章  系统架构", level=1)
    
    doc.add_heading("2.1  系统组成", level=2)
    doc.add_paragraph(
        "DKSDK 系统由以下核心组件构成："
    )
    doc.add_paragraph("MCU 主控层：负责调用 SDK 接口，注册回调函数，管理通信通道。", style='List Bullet')
    doc.add_paragraph("SDK 核心层（DKSdk.c/h）：运行在 SE 安全芯片内部，管理认证状态机、APDU 指令队列、BLE 帧解析等。", style='List Bullet')
    doc.add_paragraph("SE 硬件抽象层：通过 SE Agent 提供的 APDU 通道与 SE 硬件交互。", style='List Bullet')
    doc.add_paragraph("DK 设备层：通过 NFC 或 BLE 通道与 SDK 交互的数字钥匙设备。", style='List Bullet')
    
    doc.add_heading("2.2  SDK 与 MCU 的接口模型", level=2)
    doc.add_paragraph(
        "SDK 与 MCU 之间的接口分为两个方向："
    )
    p = doc.add_paragraph()
    run = p.add_run("SDK → MCU 接口（第4章）：")
    run.font.bold = True
    doc.add_paragraph(
        "由 SDK 提供、MCU 直接调用的函数接口。包括初始化、认证发起、数据上报、回调处理等。"
        "MCU 通过调用这些接口驱动 SDK 完成数字钥匙的各项功能。"
    )
    
    p = doc.add_paragraph()
    run = p.add_run("MCU → SDK 接口（第5章）：")
    run.font.bold = True
    doc.add_paragraph(
        "由 MCU 实现、SDK 通过函数指针回调的接口。MCU 需要将实现的回调函数通过 "
        "Sdk_Register* 系列函数注册到 SDK 中。SDK 在需要发送 APDU 指令、发送数据、"
        "通知认证结果等场景下调用这些回调函数。这些接口以 typedef 函数指针的形式定义在头文件中。"
    )
    
    doc.add_paragraph(
        "两类接口的区分规则：头文件 DKSdk.h 中以 typedef 定义的函数指针类型均为 MCU→SDK "
        "接口（即 MCU 需要实现并提供给 SDK）；其余函数声明均为 SDK→MCU 接口（即 SDK 提供给 MCU 调用）。"
    )
    
    doc.add_heading("2.3  认证流程概述", level=2)
    doc.add_paragraph(
        "DKSDK 采用状态机驱动的认证流程，主要步骤如下："
    )
    doc.add_paragraph("（1）MCU 调用 Sdk_Init() 初始化 SDK，SDK 发送 SELECT 指令选择 SE Applet；", style='List Number')
    doc.add_paragraph("（2）MCU 调用 Sdk_Auth() 发起认证，SDK 通过状态机驱动认证流程；", style='List Number')
    doc.add_paragraph("（3）SDK 依次执行：Select SE → Select DK → GPD → Internal Auth → DK Auth → External Auth；", style='List Number')
    doc.add_paragraph("（4）认证完成后，SDK 通过回调通知 MCU 认证结果；", style='List Number')
    doc.add_paragraph("（5）MCU 需周期调用 Sdk_routine() 驱动 SDK 内部状态机和超时检测。", style='List Number')
    
    # ====================================================================
    # 第3章 文件说明
    # ====================================================================
    doc.add_heading("第3章  文件说明", level=1)
    
    doc.add_heading("3.1  DKSdk.h —— SDK 头文件", level=2)
    doc.add_paragraph(
        "DKSdk.h 是 SDK 的核心头文件，定义了所有对外公开的宏、枚举、结构体和函数声明。"
        "MCU 端开发者应首先包含此头文件以使用 SDK 功能。"
    )
    
    # 3.1.1 宏定义
    doc.add_heading("3.1.1  宏定义", level=3)
    doc.add_paragraph("以下为 DKSdk.h 中定义的关键宏，MCU 可根据实际需求进行调整：")
    
    add_table_with_style(doc,
        ["宏名称", "默认值", "说明"],
        [
            ["SDK_READER_TYPE_SIZE", "6", "读卡器类型标识长度"],
            ["SDK_READER_ID_SIZE", "16", "读卡器ID长度"],
            ["SDK_READER_KEY_PARAMETER_MAX_SIZE", "16", "读卡器密钥参数最大长度"],
            ["SDK_CARD_ID_SIZE", "16", "卡片ID长度"],
            ["SDK_CHANNEL_ID_MAX_SIZE", "6", "通道ID最大长度"],
            ["SDK_TIMEOUT_OVERALL_TICKS", "30", "总认证时间限制（单位：100ms Tick），即3秒"],
            ["SDK_TIMEOUT_STEP_TICKS", "10", "单步最大耗时限制（单位：100ms Tick），即1秒"],
            ["SDK_LOG_LEVEL_NONE", "0", "关闭所有日志输出"],
            ["SDK_LOG_LEVEL_ERROR", "1", "仅开启 Error 日志（含错误和超时信息）"],
            ["SDK_LOG_LEVEL_DEBUG", "2", "开启 Debug 日志（含详细调试信息）"],
            ["SDK_LOG_LEVEL", "SDK_LOG_LEVEL_ERROR", "当前日志等级配置"],
        ],
        col_widths=[5.5, 3.5, 7.0]
    )
    
    # 错误码
    doc.add_paragraph("以下为 SDK 定义的错误码，用于认证结果通知：")
    add_table_with_style(doc,
        ["错误码", "宏名称", "说明"],
        [
            ["0x10", "ERR_AUTH_BASE", "认证错误基准值"],
            ["0x11", "ERR_AUTH_SELECT_SE_FAIL", "SE Applet 选择失败"],
            ["0x12", "ERR_AUTH_SELECT_DK_FAIL", "DK Applet 选择失败"],
            ["0x13", "ERR_AUTH_GPD_FAIL", "GPD 指令失败"],
            ["0x14", "ERR_AUTH_INTERNAL_FAIL", "内部认证失败"],
            ["0x15", "ERR_AUTH_DK_AUTH_FAIL", "DK 认证失败"],
            ["0x16", "ERR_AUTH_EXTERNAL_FAIL", "外部认证失败"],
            ["0x17", "ERR_AUTH_NO_PERMISSION", "无权限"],
            ["0x18", "ERR_AUTH_TIMEOUT", "认证超时"],
            ["0x19", "ERR_COMMAND_TIMEOUT", "指令超时"],
            ["0x1A", "ERR_NOT_INIT", "未初始化"],
            ["0x21", "ERR_PARSE_CRC_FAIL", "CRC 校验失败"],
            ["0x22", "ERR_PARSE_FRAME_TOO_SHORT", "帧长度过短"],
            ["0x23", "ERR_PARSE_SOF_WRONG", "帧起始符错误"],
            ["0x24", "ERR_PARSE_MESSAGEID_WRONG", "消息ID错误"],
            ["0x25", "ERR_PARSE_COMMANDID_WRONG", "指令ID错误"],
            ["0x26", "ERR_PARSE_TLV_WRONG", "TLV 解析错误"],
            ["0x27", "ERR_PARSE_SW_WRONG", "状态字 SW 错误"],
        ],
        col_widths=[1.5, 5.5, 9.0]
    )
    
    # 3.1.2 枚举类型
    doc.add_heading("3.1.2  枚举类型定义", level=3)
    
    doc.add_heading("Sdk_ChannelType", level=4)
    doc.add_paragraph("通道类型枚举，用于标识通信通道的物理层类型。")
    add_table_with_style(doc,
        ["枚举值", "数值", "说明"],
        [
            ["SDK_CHANNEL_TYPE_NFC", "0", "NFC 近场通信通道"],
            ["SDK_CHANNEL_TYPE_BLE", "1", "BLE 低功耗蓝牙通信通道"],
        ],
        col_widths=[5.0, 2.0, 9.0]
    )
    
    # 3.1.3 结构体定义
    doc.add_heading("3.1.3  结构体定义与使用说明", level=3)
    
    # --- Sdk_Channel ---
    doc.add_heading("Sdk_Channel（通道描述结构体）", level=4)
    doc.add_paragraph("描述一个通信通道的完整信息，包括通道类型、协议类型和通道标识。该结构体在几乎所有接口中作为通道的标识被使用。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["channel_type", "Sdk_ChannelType", "通道类型（NFC=0, BLE=1）"],
            ["protocol_type", "uint8_t", "协议类型（0x01=ICCE协议）"],
            ["id[SDK_CHANNEL_ID_MAX_SIZE]", "uint8_t[]", "通道ID数组，NFC通道 id[0]=0~1，BLE通道 id[0]=0~3"],
            ["idSize", "uint8_t", "通道ID有效长度"],
        ],
        col_widths=[4.0, 3.5, 8.5]
    )
    
    # --- Sdk_SendParam ---
    doc.add_heading("Sdk_SendParam（发送参数结构体）", level=4)
    doc.add_paragraph("用于 SDK 通过回调向 MCU 传递待发送数据的参数结构体。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["channel", "Sdk_Channel", "目标通信通道"],
            ["*dataBuffer", "uint8_t*", "待发送数据缓冲区指针"],
            ["dataSize", "uint16_t", "待发送数据长度（字节）"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_ApduParam ---
    doc.add_heading("Sdk_ApduParam（APDU指令参数结构体）", level=4)
    doc.add_paragraph("用于 SDK 通过回调向 MCU 传递 APDU 指令的参数结构体。MCU 需要将 dataBuffer 中的数据通过 SE 硬件发送出去。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["*dataBuffer", "uint8_t*", "APDU 指令数据缓冲区指针"],
            ["dataSize", "uint16_t", "APDU 指令数据长度（字节）"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_RKECallbackParam ---
    doc.add_heading("Sdk_RKECallbackParam（RKE回调参数结构体）", level=4)
    doc.add_paragraph("RKE 解密结果参数，SE 解密 RKE 密文后通过此结构体将结果返回给 MCU。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["rkeResult[8]", "uint8_t[8]", "RKE 解密结果，8字节固定长度"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_NotifyAuthResultParam ---
    doc.add_heading("Sdk_NotifyAuthResultParam（认证结果通知结构体）", level=4)
    doc.add_paragraph("认证完成后，SDK 通过此结构体向 MCU 报告认证结果。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["channel", "Sdk_Channel", "认证对应的通信通道"],
            ["errorCode", "uint8_t", "错误码（0x00=成功，其他值参考错误码表）"],
            ["sw[2]", "uint8_t[2]", "APDU 状态字 SW（SW1, SW2）"],
            ["cardId[16]", "uint8_t[16]", "认证成功时的卡片ID"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_CalibData ---
    doc.add_heading("Sdk_CalibData（校准数据通知结构体）", level=4)
    doc.add_paragraph("SDK 通过此结构体将校准数据通知给 MCU。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["channel", "Sdk_Channel", "通信通道"],
            ["*data", "uint8_t*", "校准数据缓冲区指针"],
            ["dataSize", "uint8_t", "校准数据长度（字节）"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_GetCalibReq ---
    doc.add_heading("Sdk_GetCalibReq（校准数据请求结构体）", level=4)
    doc.add_paragraph("MCU 调用 Sdk_GetCalibData 时传入此结构体，指定通道和请求参数。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["channel", "Sdk_Channel", "通信通道"],
            ["*dataBuffer", "uint8_t*", "请求数据缓冲区指针"],
            ["dataSize", "uint16_t", "请求数据长度（字节）"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_DataReportParam ---
    doc.add_heading("Sdk_DataReportParam（数据上报参数结构体）", level=4)
    doc.add_paragraph("MCU 调用 Sdk_DataReport 时传入此结构体，指定上报通道和上报数据。dataBuffer 为定长16字节。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["channel", "Sdk_Channel", "通信通道"],
            ["dataBuffer[16]", "uint8_t[16]", "上报数据缓冲区（定长16字节）"],
            ["dataSize", "uint16_t", "上报数据有效长度（≤16字节）"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_VersionParam ---
    doc.add_heading("Sdk_VersionParam（版本信息结构体）", level=4)
    doc.add_paragraph("SDK 和 Applet 的版本信息，通过版本回调通知给 MCU。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["SDK_Version[4]", "uint8_t[4]", "SDK 版本号（如 {0x26,0x04,0x29,0x01}）"],
            ["applet_version[4]", "uint8_t[4]", "Applet 版本号"],
            ["SEID[16]", "uint8_t[16]", "SE 安全芯片唯一标识"],
        ],
        col_widths=[3.5, 3.0, 9.5]
    )
    
    # --- Sdk_RKEParam ---
    doc.add_heading("Sdk_RKEParam（RKE指令参数结构体）", level=4)
    doc.add_paragraph("SDK 通过 RKE 回调将此结构体传递给 MCU，MCU 据此向车辆发送 RKE 指令。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["channel", "Sdk_Channel", "通信通道"],
            ["RKEcmd", "uint8_t", "RKE 指令码（由SE解密后的明文指令）"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- SEAgent_ApduParam ---
    doc.add_heading("SEAgent_ApduParam（SE Agent APDU参数结构体）", level=4)
    doc.add_paragraph("SE Agent 发送 APDU 指令的参数结构体，用于 TSM 等应用场景。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["*reqBuffer", "uint8_t*", "APDU 请求数据缓冲区指针"],
            ["reqSize", "uint16_t", "APDU 请求数据长度"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- SEAgent_ExecTsmParam ---
    doc.add_heading("SEAgent_ExecTsmParam（TSM执行参数结构体）", level=4)
    doc.add_paragraph("MCU 调用 SEAgent_ExecTsmCmd 时传入此结构体，指定 TSM 指令数据。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["*reqBuffer", "uint8_t*", "TSM 指令请求数据缓冲区"],
            ["reqSize", "uint16_t", "TSM 指令请求数据长度"],
        ],
        col_widths=[3.0, 3.0, 10.0]
    )
    
    # --- Sdk_Config_t ---
    doc.add_heading("Sdk_Config_t（SDK回调配置结构体）", level=4)
    doc.add_paragraph(
        "SDK 的回调配置结构体，聚合了所有 MCU 需要注册的回调函数指针。"
        "MCU 不需要直接操作此结构体，而是通过 Sdk_Register* 系列函数逐个注册回调。"
        "该结构体为 SDK 内部使用（全局变量 g_sdk_cbs），用于统一管理所有回调。"
    )
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["apdu_cb", "Sdk_Apdu", "APDU 指令发送回调"],
            ["send_cb", "Sdk_Send", "数据发送回调"],
            ["rke_cb", "Sdk_RKE", "RKE 指令回调"],
            ["notify_auth_cb", "Sdk_NotifyAuthResult", "认证结果通知回调"],
            ["notify_calib_cb", "Sdk_NotifyCalibData", "校准数据通知回调"],
            ["get_log_cb", "Sdk_GetLogCallback", "日志获取回调"],
            ["get_version_cb", "Sdk_GetVersionCallback", "版本信息回调"],
            ["data_report_cb", "Sdk_DataReportCallback", "数据上报结果回调"],
            ["seagent_apdu_cb", "SEAgent_apdu", "SE Agent APDU 回调"],
            ["seagent_apply_cb", "SEAgent_ApplySeOper", "SE 操作申请回调"],
            ["seagent_release_cb", "SEAgent_ReleaSeOper", "SE 操作释放回调"],
            ["seagent_update_cb", "SEAgent_AppletUpdate", "Applet 更新结果回调"],
        ],
        col_widths=[4.0, 4.5, 7.5]
    )
    
    # --- 3.2 DKSdk.c ---
    doc.add_heading("3.2  DKSdk.c —— SDK 实现文件", level=2)
    doc.add_paragraph(
        "DKSdk.c 是 SDK 的核心实现文件，包含认证状态机、BLE 帧解析、APDU 指令队列管理等内部逻辑。"
        "MCU 开发者通常不需要直接了解此文件中的内部实现，但以下内部数据结构有助于理解 SDK 的工作机制："
    )
    
    doc.add_heading("3.2.1  内部关键数据结构", level=3)
    
    doc.add_paragraph("以下为 DKSdk.c 中定义的内部结构体，供开发者理解 SDK 内部机制参考，不对外开放。")
    
    # ICCE_FrameInfo_t
    doc.add_heading("ICCE_FrameInfo_t（ICCE帧信息）", level=4)
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["control_field", "uint8_t", "控制字段（bit4=is_request）"],
            ["is_request", "uint8_t", "是否为请求帧"],
            ["message_id", "uint8_t", "消息ID（0x01=AUTH, 0x02=COMMAND, 0x03=NOTIFICATION）"],
            ["command_id", "uint8_t", "命令ID（0x01=AUTH, 0x02=CMD, 0x03=RKE, 0x04=GETINFO）"],
            ["fsn", "uint8_t", "帧序列号"],
            ["*payload", "const uint8_t*", "帧有效载荷数据指针"],
            ["payload_length", "uint16_t", "帧有效载荷长度"],
        ],
        col_widths=[3.5, 3.0, 9.5]
    )
    
    # Sdk_SessionContext
    doc.add_heading("Sdk_SessionContext（会话上下文）", level=4)
    doc.add_paragraph("每个通信通道对应一个会话上下文，管理该通道的认证状态、超时计数和独立缓冲区。最大支持 SDK_MAX_SESSIONS（6）个会话。")
    add_table_with_style(doc,
        ["字段名", "类型", "说明"],
        [
            ["is_active", "uint8_t", "会话是否激活"],
            ["channel", "Sdk_Channel", "通道信息"],
            ["status", "uint8_t", "认证状态机当前状态"],
            ["overall_ticks", "uint16_t", "会话总超时计数器"],
            ["peer_step_ticks", "uint16_t", "等待对端回调超时计数器"],
            ["is_waiting_peer", "uint8_t", "是否正在等待对端回调"],
            ["card_id[16]", "uint8_t[16]", "卡片ID"],
            ["status_word", "uint16_t", "状态字 SW"],
            ["apdu_resp[]", "uint8_t[260]", "SE APDU 响应缓冲区（每会话独立）"],
            ["dk_apdu_resp[]", "uint8_t[260]", "DK/BLE 响应缓冲区（每会话独立）"],
            ["reader_rnd[8]", "uint8_t[8]", "读卡器随机数"],
            ["reader_key_parameter[16]", "uint8_t[16]", "读卡器密钥参数"],
            ["has_pending_rke", "uint8_t", "是否有挂起的 RKE 请求"],
        ],
        col_widths=[4.0, 3.0, 9.0]
    )
    
    # Sdk_HwApduManager
    doc.add_heading("Sdk_HwApduManager（APDU硬件命令管理器）", level=4)
    doc.add_paragraph("管理 SE APDU 指令队列，实现多会话并发时的指令排队和超时控制。最大队列深度为 SDK_MAX_HW_QUEUE（6）。使用队列锁（queue_locked）保护并发操作。")
    
    # 状态机状态常量
    doc.add_heading("3.2.2  状态机状态常量", level=3)
    add_table_with_style(doc,
        ["状态宏", "值", "说明"],
        [
            ["STATE_MACHINE_INIT", "0x00", "初始化状态"],
            ["STATE_MACHINE_SELECT_SE", "0x01", "选择 SE Applet"],
            ["STATE_MACHINE_SELECT_DK", "0x02", "选择 DK Applet"],
            ["STATE_MACHINE_GPD", "0x03", "获取 GPD 数据"],
            ["STATE_MACHINE_INTERNAL_AUTH", "0x04", "内部认证"],
            ["STATE_MACHINE_AUTH", "0x05", "DK 认证"],
            ["STATE_MACHINE_EXTERNAL_AUTH", "0x06", "外部认证"],
            ["STATE_MACHINE_DATAREPORT_REQUEST", "0x08", "数据上报请求"],
            ["STATE_MACHINE_DATAREPORT_RESPONSE", "0x09", "数据上报响应"],
            ["STATE_MACHINE_RKE_REQUEST", "0x0A", "RKE 请求"],
            ["STATE_MACHINE_RKE_RESPONSE", "0x0B", "RKE 响应"],
            ["STATE_MACHINE_CALIBDATA", "0x0E", "校准数据"],
            ["STATE_MACHINE_OFFLINE_AUTH", "0x0F", "离线认证"],
        ],
        col_widths=[6.0, 2.0, 8.0]
    )
    
    # ====================================================================
    # 第4章 SDK 提供给 MCU 的接口
    # ====================================================================
    doc.add_heading("第4章  SDK 提供给 MCU 的接口", level=1)
    doc.add_paragraph(
        "本章描述 SDK 提供给 MCU 直接调用的所有接口函数。MCU 通过调用这些函数驱动 SDK 完成数字钥匙的各项功能。"
        "接口按功能分为以下几类：初始化与调度、认证管理、回调处理、数据操作、注册函数和 SE Agent 接口。"
    )
    
    # 第4章总体接口汇总表
    doc.add_paragraph("表4-0  SDK→MCU 全部接口总览")
    add_interface_summary_table(doc,
        ["分类", "函数名", "函数原型", "功能描述"],
        [
            ["初始化与调度", "Sdk_Init", "void Sdk_Init(void)", "初始化SDK，自动发送SE Applet选择指令"],
            ["初始化与调度", "Sdk_routine", "void Sdk_routine(void)", "SDK主循环：超时检测+队列驱动+响应分发"],
            ["认证管理", "Sdk_Auth", "uint8_t Sdk_Auth(Sdk_Channel *nChannel)", "发起数字钥匙认证"],
            ["认证管理", "Sdk_CancelAuth", "uint8_t Sdk_CancelAuth(Sdk_Channel *nChannel)", "取消认证（NULL=全部重置）"],
            ["认证管理", "Sdk_Release", "void Sdk_Release(Sdk_Channel *nChannel)", "释放通道会话"],
            ["回调处理", "Sdk_ApduCallback", "uint8_t Sdk_ApduCallback(uint8_t, uint8_t*, uint16_t)", "APDU响应回传（ISR安全）"],
            ["回调处理", "Sdk_SendCallback", "uint8_t Sdk_SendCallback(Sdk_Channel*, uint8_t, uint8_t*, uint16_t)", "DK响应回传+状态机分发"],
            ["回调处理", "Sdk_RKECallback", "uint8_t Sdk_RKECallback(Sdk_Channel*, Sdk_RKECallbackParam*)", "RKE执行结果回传"],
            ["数据操作", "Sdk_DataReport", "uint8_t Sdk_DataReport(Sdk_DataReportParam*)", "数据上报（加密后发送）"],
            ["数据操作", "Sdk_GetCalibData", "uint8_t Sdk_GetCalibData(Sdk_GetCalibReq*)", "获取校准数据"],
            ["数据操作", "Sdk_GetLog", "uint8_t Sdk_GetLog(uint16_t)", "获取SDK日志"],
            ["SE Agent", "SEAgent_ApduCallback", "uint8_t SEAgent_ApduCallback(uint8_t, uint8_t*, uint16_t)", "SE Agent APDU响应（预留）"],
            ["SE Agent", "SEAgent_ExecTsmCmd", "uint8_t SEAgent_ExecTsmCmd(SEAgent_ExecTsmParam*)", "执行TSM指令（预留）"],
            ["注册函数", "Sdk_RegisterApdu", "void Sdk_RegisterApdu(Sdk_Apdu)", "注册APDU发送回调"],
            ["注册函数", "Sdk_RegisterSend", "void Sdk_RegisterSend(Sdk_Send)", "注册数据发送回调"],
            ["注册函数", "Sdk_RegisterRKE", "void Sdk_RegisterRKE(Sdk_RKE)", "注册RKE指令回调"],
            ["注册函数", "Sdk_RegisterNotifyAuthResult", "void Sdk_RegisterNotifyAuthResult(Sdk_NotifyAuthResult)", "注册认证结果通知回调"],
            ["注册函数", "Sdk_RegisterGetLogCallback", "void Sdk_RegisterGetLogCallback(Sdk_GetLogCallback)", "注册日志输出回调"],
            ["注册函数", "Sdk_RegisterGetVersionCallback", "void Sdk_RegisterGetVersionCallback(Sdk_GetVersionCallback)", "注册版本信息回调"],
            ["注册函数", "Sdk_RegisterNotifyCalibData", "void Sdk_RegisterNotifyCalibData(Sdk_NotifyCalibData)", "注册校准数据通知回调"],
            ["注册函数", "Sdk_RegisterDataReportCallback", "void Sdk_RegisterDataReportCallback(Sdk_DataReportCallback)", "注册数据上报结果回调"],
        ],
        col_widths=[2.0, 3.0, 5.5, 4.5]
    )
    doc.add_paragraph("以下各节对各接口进行详细说明。")
    
    # ---- 4.1 初始化与调度 ----
    doc.add_heading("4.1  初始化与调度", level=2)
    
    # 接口汇总表
    doc.add_paragraph("表4-1  初始化与调度接口汇总")
    add_interface_summary_table(doc,
        ["函数名", "函数原型", "功能描述", "返回值"],
        [
            ["Sdk_Init", "void Sdk_Init(void)", "初始化SDK，清零会话/队列，自动发送SE Applet选择指令", "无"],
            ["Sdk_routine", "void Sdk_routine(void)", "SDK主循环调度，检查超时、驱动队列、处理APDU响应", "无"],
        ],
        col_widths=[2.5, 4.5, 6.5, 1.5]
    )
    
    add_function_detail(doc,
        "Sdk_Init",
        "void Sdk_Init(void);",
        "初始化 SDK 内部状态。清零所有会话上下文和硬件命令管理器，并自动发送第一帧 SELECT 指令选择 SE Applet。",
        [("无", "void", "—")],
        "无返回值。",
        "应在系统上电后、调用任何其他 SDK 接口之前调用。初始化完成后，MCU 应周期性调用 Sdk_routine() 驱动 SDK。"
    )
    
    add_function_detail(doc,
        "Sdk_routine",
        "void Sdk_routine(void);",
        "SDK 主循环调度函数。负责检查所有会话的超时状态、驱动 SE 指令队列、处理 ISR 缓存的 APDU 响应。",
        [("无", "void", "—")],
        "无返回值。",
        "MCU 应以固定周期（建议 1ms）调用此函数。此函数执行超时检测：会话总超时 3s、SE 单指令超时 1s、等待对端回调超时 1s。超时后将自动结束会话并通过回调通知 MCU。"
    )
    
    # ---- 4.2 认证管理 ----
    doc.add_heading("4.2  认证管理", level=2)
    
    # 接口汇总表
    doc.add_paragraph("表4-2  认证管理接口汇总")
    add_interface_summary_table(doc,
        ["函数名", "函数原型", "功能描述", "返回值"],
        [
            ["Sdk_Auth", "uint8_t Sdk_Auth(Sdk_Channel *nChannel)", "对指定通道发起数字钥匙认证流程", "0=成功, 1=失败"],
            ["Sdk_CancelAuth", "uint8_t Sdk_CancelAuth(Sdk_Channel *nChannel)", "取消指定通道认证（nChannel=NULL则重置全部）", "0=成功"],
            ["Sdk_Release", "void Sdk_Release(Sdk_Channel *nChannel)", "释放指定通道会话资源", "无"],
        ],
        col_widths=[2.5, 5.5, 5.5, 1.5]
    )
    
    add_function_detail(doc,
        "Sdk_Auth",
        "uint8_t Sdk_Auth(Sdk_Channel *nChannel);",
        "对指定通道发起数字钥匙认证流程。SDK 将自动驱动状态机完成：Select SE → GPD → Internal Auth → DK Auth → External Auth 全流程。认证结果通过 notify_auth_cb 回调通知 MCU。",
        [
            ("nChannel", "Sdk_Channel*", "指向通道描述结构体的指针，指定认证使用的通信通道"),
        ],
        "STATUS_SUCCESS（0）：认证流程已启动。\nSTATUS_FAILED（1）：参数无效或通道正忙（被 RKE/DataReport/GetCalibData 占用）。",
        "（1）若通道正在进行 RKE、DataReport 或 GetCalibData 操作，认证将被阻塞，MCU 应等待当前操作完成后再重试。\n"
        "（2）认证结果为异步通知，MCU 需通过注册的 notify_auth_cb 回调接收结果。"
    )
    
    add_function_detail(doc,
        "Sdk_CancelAuth",
        "uint8_t Sdk_CancelAuth(Sdk_Channel *nChannel);",
        "取消指定通道上的认证流程。若 nChannel 为 NULL，则重置所有会话（等价于重新初始化）。结束会话前会自动处理该通道上挂起的 RKE 请求，确保不丢失数据。",
        [
            ("nChannel", "Sdk_Channel*", "指向通道描述结构体的指针。若为 NULL，则重置全部会话。"),
        ],
        "STATUS_SUCCESS（0）：操作成功。",
        "内部通过 Sdk_FinishSession 结束会话，会自动处理 pending RKE 数据。"
    )
    
    add_function_detail(doc,
        "Sdk_Release",
        "void Sdk_Release(Sdk_Channel *nChannel);",
        "释放指定通道的会话资源。内部调用 Sdk_CancelAuth 完成释放。",
        [
            ("nChannel", "Sdk_Channel*", "指向通道描述结构体的指针"),
        ],
        "无返回值。",
        "与 Sdk_CancelAuth 功能等价。"
    )
    
    # ---- 4.3 回调处理 ----
    doc.add_heading("4.3  回调处理（ISR安全）", level=2)
    doc.add_paragraph(
        "以下回调处理函数由 MCU 在接收到 SE 或 DK 的响应数据后调用，将数据回传给 SDK。"
        "这些函数设计为 ISR 安全：Sdk_ApduCallback 仅将数据拷贝到缓冲区，实际处理由 Sdk_routine 主循环完成。"
    )
    
    # 接口汇总表
    doc.add_paragraph("表4-3  回调处理接口汇总")
    add_interface_summary_table(doc,
        ["函数名", "函数原型", "功能描述", "返回值"],
        [
            ["Sdk_ApduCallback", "uint8_t Sdk_ApduCallback(uint8_t result, uint8_t *dataBuffer, uint16_t dataSize)",
             "APDU响应回调（ISR安全），将SE响应数据缓存到转发缓冲区", "0=成功, 1=overflow"],
            ["Sdk_SendCallback", "uint8_t Sdk_SendCallback(Sdk_Channel *nChannel, uint8_t result, uint8_t *dataBuffer, uint16_t dataSize)",
             "DK数据发送响应回调，根据状态机分发到认证/RKE/上报流程", "0=成功, 1=失败"],
            ["Sdk_RKECallback", "uint8_t Sdk_RKECallback(Sdk_Channel *nChannel, Sdk_RKECallbackParam *param)",
             "RKE指令执行结果回调，将MCU执行结果加密后发给DK", "0=成功, 1=失败"],
        ],
        col_widths=[2.3, 6.2, 5.0, 1.5]
    )
    
    add_function_detail(doc,
        "Sdk_ApduCallback",
        "uint8_t Sdk_ApduCallback(uint8_t result, uint8_t *dataBuffer, uint16_t dataSize);",
        "APDU 指令响应回调。MCU 在 SE 返回 APDU 响应后调用此函数，将响应数据回传给 SDK。该函数设计为 ISR 安全，仅将数据拷贝到 ISR→主循环转发缓冲区，实际队列操作和业务处理由主循环 Sdk_ProcessApduResponse 完成。",
        [
            ("result", "uint8_t", "APDU 执行结果（0x00=成功）"),
            ("dataBuffer", "uint8_t*", "APDU 响应数据缓冲区指针"),
            ("dataSize", "uint16_t", "APDU 响应数据长度"),
        ],
        "STATUS_SUCCESS（0）：数据已缓存。\nSTATUS_FAILED（1）：上一帧响应尚未被主循环消费（overflow）。",
        "（1）可在 ISR 中调用，但必须保证不产生 overflow。\n（2）若返回 STATUS_FAILED，说明主循环消费速度不够，上一帧数据未被取走。"
    )
    
    add_function_detail(doc,
        "Sdk_SendCallback",
        "uint8_t Sdk_SendCallback(Sdk_Channel *nChannel, uint8_t result, uint8_t *dataBuffer, uint16_t dataSize);",
        "DK 数据发送响应回调。MCU 在收到 DK 设备的响应数据后调用此函数，SDK 将根据当前状态机状态决定下一步操作（继续认证流程、执行 RKE 解密、处理数据上报等）。",
        [
            ("nChannel", "Sdk_Channel*", "通道信息"),
            ("result", "uint8_t", "发送结果（0x00=成功）"),
            ("dataBuffer", "uint8_t*", "DK 响应数据缓冲区"),
            ("dataSize", "uint16_t", "响应数据长度"),
        ],
        "STATUS_SUCCESS（0）：数据已处理。\nSTATUS_FAILED（1）：处理失败。",
        "（1）BLE 通道的数据会先经过 ICCE 帧解析（icce_parse_frame）提取 payload。\n"
        "（2）NFC 通道的数据直接解析 SW 状态字后进行状态机分发。\n"
        "（3）若收到 RKE 数据但通道正忙，RKE 数据将暂存到 pending 队列，待当前操作完成后自动处理。"
    )
    
    add_function_detail(doc,
        "Sdk_RKECallback",
        "uint8_t Sdk_RKECallback(Sdk_Channel *nChannel, Sdk_RKECallbackParam *param);",
        "RKE 指令执行结果回调。MCU 在执行 RKE 指令后调用此函数，将执行结果回传给 SDK。SDK 会将结果加密后发送给 DK 设备。",
        [
            ("nChannel", "Sdk_Channel*", "通道信息"),
            ("param", "Sdk_RKECallbackParam*", "RKE 回调参数，包含 8 字节的 rkeResult"),
        ],
        "STATUS_SUCCESS（0）：处理成功。\nSTATUS_FAILED（1）：参数无效或未找到匹配会话。",
        "MCU 应在执行 RKE 指令后尽快调用此函数，将执行结果（如车门状态）返回给 SDK。"
    )
    
    # ---- 4.4 数据操作 ----
    doc.add_heading("4.4  数据操作", level=2)
    
    # 接口汇总表
    doc.add_paragraph("表4-4  数据操作接口汇总")
    add_interface_summary_table(doc,
        ["函数名", "函数原型", "功能描述", "返回值"],
        [
            ["Sdk_DataReport", "uint8_t Sdk_DataReport(Sdk_DataReportParam *param)",
             "向DK设备上报数据（车辆状态等），经SE加密后发送", "0=成功, 1=失败"],
            ["Sdk_GetCalibData", "uint8_t Sdk_GetCalibData(Sdk_GetCalibReq *req)",
             "向SE请求校准数据，结果通过notify_calib_cb回调通知", "0=成功, 1=失败"],
            ["Sdk_GetLog", "uint8_t Sdk_GetLog(uint16_t bufferMaxSize)",
             "获取SDK日志（通过get_log_cb回调输出，等级由SDK_LOG_LEVEL控制）", "0=成功"],
        ],
        col_widths=[2.5, 5.5, 5.5, 1.5]
    )
    
    add_function_detail(doc,
        "Sdk_DataReport",
        "uint8_t Sdk_DataReport(Sdk_DataReportParam *param);",
        "数据上报接口。MCU 通过此接口向 DK 设备上报数据（如车辆状态信息）。数据经 SE 加密后通过通信通道发送给 DK 设备。",
        [
            ("param", "Sdk_DataReportParam*", "数据上报参数，包含通道信息、16字节数据缓冲区及有效长度"),
        ],
        "STATUS_SUCCESS（0）：上报流程已启动。\nSTATUS_FAILED（1）：参数无效或通道正忙。",
        "（1）若通道正被 RKE/Auth/GetCalibData 占用，上报将被阻塞，MCU 应在 2s 后重试。\n"
        "（2）上报数据最大有效长度为 16 字节（dataBuffer 定长）。\n"
        "（3）上报结果通过 data_report_cb 回调通知 MCU。"
    )
    
    add_function_detail(doc,
        "Sdk_GetCalibData",
        "uint8_t Sdk_GetCalibData(Sdk_GetCalibReq *req);",
        "获取校准数据接口。MCU 通过此接口向 SE 请求校准数据，SE 通过加密通道获取后通过 notify_calib_cb 回调通知 MCU。",
        [
            ("req", "Sdk_GetCalibReq*", "校准数据请求参数，包含通道信息和请求数据"),
        ],
        "STATUS_SUCCESS（0）：请求已发送。\nSTATUS_FAILED（1）：参数无效或通道正忙。",
        "若通道正被 RKE/Auth/DataReport 占用，请求将被阻塞。"
    )
    
    add_function_detail(doc,
        "Sdk_GetLog",
        "uint8_t Sdk_GetLog(uint16_t bufferMaxSize);",
        "获取 SDK 日志接口。MCU 通过注册的 get_log_cb 回调接收日志数据。实际的日志输出由 SDK 内部的 Sdk_LogError/Sdk_LogDebug 函数通过 get_log_cb 回调完成。",
        [
            ("bufferMaxSize", "uint16_t", "日志缓冲区最大大小（预留参数）"),
        ],
        "STATUS_SUCCESS（0）",
        "日志输出等级由宏 SDK_LOG_LEVEL 控制（默认仅输出 Error 日志）。"
    )
    
    # ---- 4.5 SE Agent 接口 ----
    doc.add_heading("4.5  SE Agent 接口", level=2)
    
    # 接口汇总表
    doc.add_paragraph("表4-5  SE Agent 接口汇总")
    add_interface_summary_table(doc,
        ["函数名", "函数原型", "功能描述", "返回值"],
        [
            ["SEAgent_ApduCallback", "uint8_t SEAgent_ApduCallback(uint8_t result, uint8_t *respData, uint16_t dataSize)",
             "SE Agent APDU响应回调（预留接口）", "0=成功"],
            ["SEAgent_ExecTsmCmd", "uint8_t SEAgent_ExecTsmCmd(SEAgent_ExecTsmParam *param)",
             "执行TSM指令，用于应用更新/个人化（预留接口）", "0=成功"],
        ],
        col_widths=[3.0, 6.5, 4.0, 1.5]
    )
    
    add_function_detail(doc,
        "SEAgent_ApduCallback",
        "uint8_t SEAgent_ApduCallback(uint8_t result, uint8_t *respData, uint16_t dataSize);",
        "SE Agent APDU 指令响应回调。用于 TSM 等应用场景中 SE Agent 的 APDU 响应处理。（当前为预留接口，参数尚未使用）",
        [
            ("result", "uint8_t", "APDU 执行结果"),
            ("respData", "uint8_t*", "响应数据缓冲区"),
            ("dataSize", "uint16_t", "响应数据长度"),
        ],
        "STATUS_SUCCESS（0）",
        "当前为预留接口，参数尚未在内部处理。"
    )
    
    add_function_detail(doc,
        "SEAgent_ExecTsmCmd",
        "uint8_t SEAgent_ExecTsmCmd(SEAgent_ExecTsmParam *param);",
        "执行 TSM（可信服务管理）指令。用于应用更新、个人化等 TSM 操作。（当前为预留接口，参数尚未使用）",
        [
            ("param", "SEAgent_ExecTsmParam*", "TSM 指令参数"),
        ],
        "STATUS_SUCCESS（0）",
        "当前为预留接口，功能尚未实现。"
    )
    
    # ---- 4.6 注册函数 ----
    doc.add_heading("4.6  回调注册函数", level=2)
    doc.add_paragraph(
        "以下注册函数用于 MCU 将实现的回调函数注册到 SDK 中。所有注册函数遵循相同的模式："
        "传入函数指针，SDK 内部存储到全局回调配置结构体 g_sdk_cbs 中。注册应在 Sdk_Init 之后、"
        "开始认证或数据操作之前完成。"
    )
    
    # 接口汇总表
    doc.add_paragraph("表4-6  回调注册函数汇总")
    add_interface_summary_table(doc,
        ["函数名", "函数原型", "注册的回调类型", "功能描述"],
        [
            ["Sdk_RegisterApdu", "void Sdk_RegisterApdu(Sdk_Apdu apduFunc)", "Sdk_Apdu",
             "注册APDU发送回调，SDK通过此回调向SE发送APDU指令"],
            ["Sdk_RegisterSend", "void Sdk_RegisterSend(Sdk_Send sendFunc)", "Sdk_Send",
             "注册数据发送回调，SDK通过此回调向DK发送数据"],
            ["Sdk_RegisterRKE", "void Sdk_RegisterRKE(Sdk_RKE rkeFunc)", "Sdk_RKE",
             "注册RKE指令回调，SE解密后SDK通过此回调传递明文指令给MCU"],
            ["Sdk_RegisterNotifyAuthResult", "void Sdk_RegisterNotifyAuthResult(Sdk_NotifyAuthResult func)", "Sdk_NotifyAuthResult",
             "注册认证结果通知回调，认证完成后通知MCU"],
            ["Sdk_RegisterGetLogCallback", "void Sdk_RegisterGetLogCallback(Sdk_GetLogCallback func)", "Sdk_GetLogCallback",
             "注册日志输出回调，SDK通过此回调输出运行日志"],
            ["Sdk_RegisterGetVersionCallback", "void Sdk_RegisterGetVersionCallback(Sdk_GetVersionCallback func)", "Sdk_GetVersionCallback",
             "注册版本信息回调，初始化后通知MCU版本和SEID"],
            ["Sdk_RegisterNotifyCalibData", "void Sdk_RegisterNotifyCalibData(Sdk_NotifyCalibData func)", "Sdk_NotifyCalibData",
             "注册校准数据通知回调，校准数据就绪后通知MCU"],
            ["Sdk_RegisterDataReportCallback", "void Sdk_RegisterDataReportCallback(Sdk_DataReportCallback func)", "Sdk_DataReportCallback",
             "注册数据上报结果回调，上报完成后通知MCU"],
        ],
        col_widths=[3.0, 5.0, 2.5, 4.5]
    )
    
    register_funcs = [
        ("Sdk_RegisterApdu", "void Sdk_RegisterApdu(Sdk_Apdu apduFunc);",
         "注册 APDU 发送回调函数。SDK 通过此回调将 APDU 指令发送给 SE 硬件执行。",
         [("apduFunc", "Sdk_Apdu", "APDU 发送回调函数指针")]),
        ("Sdk_RegisterSend", "void Sdk_RegisterSend(Sdk_Send sendFunc);",
         "注册数据发送回调函数。SDK 通过此回调将数据通过指定通道发送给 DK 设备。",
         [("sendFunc", "Sdk_Send", "数据发送回调函数指针")]),
        ("Sdk_RegisterRKE", "void Sdk_RegisterRKE(Sdk_RKE rkeFunc);",
         "注册 RKE 指令回调函数。SE 解密 RKE 密文后，SDK 通过此回调将明文指令传递给 MCU，MCU 据此执行车辆控制。",
         [("rkeFunc", "Sdk_RKE", "RKE 指令回调函数指针")]),
        ("Sdk_RegisterNotifyAuthResult", "void Sdk_RegisterNotifyAuthResult(Sdk_NotifyAuthResult notifyAuthResultFunc);",
         "注册认证结果通知回调函数。认证流程完成后，SDK 通过此回调将结果（成功/失败/错误码）通知 MCU。",
         [("notifyAuthResultFunc", "Sdk_NotifyAuthResult", "认证结果通知回调函数指针")]),
        ("Sdk_RegisterGetLogCallback", "void Sdk_RegisterGetLogCallback(Sdk_GetLogCallback getLogFunc);",
         "注册日志获取回调函数。SDK 通过此回调将运行日志输出给 MCU。日志等级由宏 SDK_LOG_LEVEL 控制。",
         [("getLogFunc", "Sdk_GetLogCallback", "日志获取回调函数指针")]),
        ("Sdk_RegisterGetVersionCallback", "void Sdk_RegisterGetVersionCallback(Sdk_GetVersionCallback getVersionFunc);",
         "注册版本信息回调函数。SDK 初始化完成后通过此回调将 SDK 版本、Applet 版本和 SEID 通知 MCU。",
         [("getVersionFunc", "Sdk_GetVersionCallback", "版本信息回调函数指针")]),
        ("Sdk_RegisterNotifyCalibData", "void Sdk_RegisterNotifyCalibData(Sdk_NotifyCalibData notifyCalibFunc);",
         "注册校准数据通知回调函数。校准数据获取完成后，SDK 通过此回调将数据通知 MCU。",
         [("notifyCalibFunc", "Sdk_NotifyCalibData", "校准数据通知回调函数指针")]),
        ("Sdk_RegisterDataReportCallback", "void Sdk_RegisterDataReportCallback(Sdk_DataReportCallback dataReportFunc);",
         "注册数据上报结果回调函数。数据上报完成后，SDK 通过此回调将上报结果通知 MCU。",
         [("dataReportFunc", "Sdk_DataReportCallback", "数据上报结果回调函数指针")]),
    ]
    
    for name, proto, desc, params in register_funcs:
        add_function_detail(doc, name, proto, desc, params,
            "无返回值。",
            "传入 NULL 将不会注册（函数内部有判空保护）。"
        )
    
    # ====================================================================
    # 第5章 MCU 提供给 SDK 的接口
    # ====================================================================
    doc.add_heading("第5章  MCU 提供给 SDK 的接口", level=1)
    doc.add_paragraph(
        "本章描述 MCU 需要实现并提供给 SDK 的回调函数接口。这些接口以 typedef 函数指针的形式定义在 "
        "DKSdk.h 中。MCU 需要按照定义的函数签名实现这些回调函数，然后通过第4.6节的注册函数注册到 SDK 中。"
    )
    doc.add_paragraph(
        "SDK 在以下场景会调用这些回调函数：发送 APDU 指令到 SE 硬件、通过 BLE/NFC 发送数据到 DK 设备、"
        "通知认证结果、传递 RKE 指令、输出日志等。MCU 应确保回调函数的实现是线程安全的，"
        "且不会长时间阻塞（特别是 apdu_cb 和 send_cb）。"
    )
    
    # 第5章总体接口汇总表
    doc.add_paragraph("表5-0  MCU→SDK 全部回调接口总览")
    add_interface_summary_table(doc,
        ["分类", "回调类型（typedef）", "函数签名", "对应注册函数"],
        [
            ["核心通信", "Sdk_Apdu", "uint8_t (*)(Sdk_ApduParam *param)", "Sdk_RegisterApdu"],
            ["核心通信", "Sdk_Send", "uint8_t (*)(Sdk_SendParam *param)", "Sdk_RegisterSend"],
            ["核心通信", "Sdk_RKE", "uint8_t (*)(Sdk_RKEParam *param)", "Sdk_RegisterRKE"],
            ["通知类", "Sdk_NotifyAuthResult", "uint8_t (*)(Sdk_NotifyAuthResultParam *param)", "Sdk_RegisterNotifyAuthResult"],
            ["通知类", "Sdk_GetVersionCallback", "uint8_t (*)(Sdk_VersionParam *param)", "Sdk_RegisterGetVersionCallback"],
            ["通知类", "Sdk_NotifyCalibData", "uint8_t (*)(Sdk_CalibData *param)", "Sdk_RegisterNotifyCalibData"],
            ["通知类", "Sdk_DataReportCallback", "uint8_t (*)(Sdk_Channel*, uint8_t*, uint16_t)", "Sdk_RegisterDataReportCallback"],
            ["通知类", "Sdk_GetLogCallback", "void (*)(uint8_t *logBuffer, uint16_t logSize)", "Sdk_RegisterGetLogCallback"],
            ["SE Agent(预留)", "SEAgent_apdu", "uint8_t (*)(SEAgent_ApduParam *param)", "（内部使用）"],
            ["SE Agent(预留)", "SEAgent_ApplySeOper", "uint8_t (*)(uint16_t duration)", "（内部使用）"],
            ["SE Agent(预留)", "SEAgent_ReleaSeOper", "uint8_t (*)(void)", "（内部使用）"],
            ["SE Agent(预留)", "SEAgent_AppletUpdate", "uint8_t (*)(uint8_t result)", "（内部使用）"],
        ],
        col_widths=[2.0, 3.5, 5.5, 4.0]
    )
    doc.add_paragraph("以下各节对各回调接口进行详细说明。")
    
    doc.add_heading("5.1  核心通信回调", level=2)
    
    # 接口汇总表
    doc.add_paragraph("表5-1  核心通信回调接口汇总")
    add_interface_summary_table(doc,
        ["回调类型（typedef）", "函数签名", "功能描述", "返回值"],
        [
            ["Sdk_Apdu", "uint8_t (*)(Sdk_ApduParam *param)",
             "APDU指令发送回调，MCU将APDU通过SE硬件发送，响应经Sdk_ApduCallback回传", "0=成功"],
            ["Sdk_Send", "uint8_t (*)(Sdk_SendParam *param)",
             "数据发送回调，MCU将数据通过NFC/BLE发送给DK，响应经Sdk_SendCallback回传", "0=成功"],
            ["Sdk_RKE", "uint8_t (*)(Sdk_RKEParam *param)",
             "RKE指令执行回调，MCU根据RKEcmd执行车辆控制，结果经Sdk_RKECallback回传", "0=成功"],
        ],
        col_widths=[2.5, 5.0, 6.0, 1.5]
    )
    
    doc.add_paragraph("以下为各回调接口的详细说明：")
    
    mcu_callbacks = [
        {
            "name": "Sdk_Apdu",
            "proto": "typedef uint8_t (*Sdk_Apdu)(Sdk_ApduParam *param);",
            "desc": "APDU 指令发送回调。MCU 实现此回调，将 SDK 传来的 APDU 指令通过 SE 硬件接口发送给安全芯片，并等待响应。响应数据通过 Sdk_ApduCallback 回传给 SDK。",
            "params": [("param", "Sdk_ApduParam*", "APDU 指令参数（含数据缓冲区指针和长度）")],
            "return": "uint8_t：0x00=成功，其他=失败",
            "notes": "此回调在 SDK 内部被调用，MCU 应尽快完成 SE 通信并调用 Sdk_ApduCallback 返回结果。建议使用异步机制（如 DMA+ISR）避免阻塞。"
        },
        {
            "name": "Sdk_Send",
            "proto": "typedef uint8_t (*Sdk_Send)(Sdk_SendParam *param);",
            "desc": "数据发送回调。MCU 实现此回调，将 SDK 传来的数据通过指定通道（NFC 或 BLE）发送给 DK 设备。响应数据通过 Sdk_SendCallback 回传给 SDK。",
            "params": [("param", "Sdk_SendParam*", "发送参数（含通道信息和数据缓冲区指针）")],
            "return": "uint8_t：0x00=成功，其他=失败",
            "notes": "MCU 需要根据 param->channel.channel_type 判断使用 NFC 还是 BLE 发送数据。BLE 通道需额外处理 ICCE 帧协议。"
        },
        {
            "name": "Sdk_RKE",
            "proto": "typedef uint8_t (*Sdk_RKE)(Sdk_RKEParam *param);",
            "desc": "RKE 指令执行回调。SE 解密 RKE 密文后，SDK 通过此回调将明文指令传递给 MCU。MCU 实现此回调，根据 RKEcmd 执行相应的车辆控制操作（如解锁车门），并将执行结果通过 Sdk_RKECallback 回传给 SDK。",
            "params": [("param", "Sdk_RKEParam*", "RKE 参数（含通道信息和 RKE 指令码）")],
            "return": "uint8_t：0x00=成功，其他=失败",
            "notes": "MCU 应在执行 RKE 指令后尽快通过 Sdk_RKECallback 返回结果，以便 SDK 将加密结果发送给 DK 设备。"
        },
    ]
    
    for cb in mcu_callbacks:
        add_callback_detail_table(doc,
            cb["name"], cb["proto"], cb["desc"], cb["params"],
            cb["return"], cb.get("notes"))
    
    doc.add_heading("5.2  通知类回调", level=2)
    
    # 接口汇总表
    doc.add_paragraph("表5-2  通知类回调接口汇总")
    add_interface_summary_table(doc,
        ["回调类型（typedef）", "函数签名", "功能描述", "返回值"],
        [
            ["Sdk_NotifyAuthResult", "uint8_t (*)(Sdk_NotifyAuthResultParam *param)",
             "认证结果通知，errorCode=0成功，非0参考错误码表；成功时cardId含卡片ID", "0=成功"],
            ["Sdk_GetVersionCallback", "uint8_t (*)(Sdk_VersionParam *param)",
             "版本信息通知，含SDK版本、Applet版本、SEID", "0=成功"],
            ["Sdk_NotifyCalibData", "uint8_t (*)(Sdk_CalibData *param)",
             "校准数据通知，MCU需在回调内完成数据拷贝", "0=成功"],
            ["Sdk_DataReportCallback", "uint8_t (*)(Sdk_Channel *nChannel, uint8_t *dataBuffer, uint16_t dataSize)",
             "数据上报结果通知，返回DK设备的响应数据", "0=成功"],
            ["Sdk_GetLogCallback", "void (*)(uint8_t *logBuffer, uint16_t logSize)",
             "日志输出回调，等级由SDK_LOG_LEVEL控制；NONE时不输出", "无"],
        ],
        col_widths=[3.0, 5.5, 5.0, 1.5]
    )
    
    doc.add_paragraph("以下为各回调接口的详细说明：")
    
    notify_callbacks = [
        {
            "name": "Sdk_NotifyAuthResult",
            "proto": "typedef uint8_t (*Sdk_NotifyAuthResult)(Sdk_NotifyAuthResultParam *param);",
            "desc": "认证结果通知回调。认证流程（成功或失败）完成后，SDK 通过此回调将结果通知 MCU。errorCode 为 0x00 表示认证成功，其他值对应具体错误（见第3.1.1节错误码表）。",
            "params": [("param", "Sdk_NotifyAuthResultParam*", "认证结果参数")],
            "return": "uint8_t：0x00=成功",
            "notes": "认证成功时 param->cardId 包含认证通过的卡片 ID（16字节）。认证失败时 param->errorCode 指示失败原因，param->sw 包含 SE 返回的状态字。"
        },
        {
            "name": "Sdk_GetVersionCallback",
            "proto": "typedef uint8_t (*Sdk_GetVersionCallback)(Sdk_VersionParam *param);",
            "desc": "版本信息回调。SDK 初始化完成后通过此回调将 SDK 版本、Applet 版本和 SE 唯一标识（SEID）通知 MCU。",
            "params": [("param", "Sdk_VersionParam*", "版本信息参数")],
            "return": "uint8_t：0x00=成功",
            "notes": "此回调在 Sdk_Init 过程中被调用。MCU 可在此回调中保存版本信息用于诊断和日志。"
        },
        {
            "name": "Sdk_NotifyCalibData",
            "proto": "typedef uint8_t (*Sdk_NotifyCalibData)(Sdk_CalibData *param);",
            "desc": "校准数据通知回调。MCU 调用 Sdk_GetCalibData 后，SDK 获取到校准数据时通过此回调通知 MCU。",
            "params": [("param", "Sdk_CalibData*", "校准数据参数")],
            "return": "uint8_t：0x00=成功",
            "notes": "回调中 param->data 指向的缓冲区在回调返回后可能被 SDK 重用，MCU 应在回调内部完成数据拷贝。"
        },
        {
            "name": "Sdk_DataReportCallback",
            "proto": "typedef uint8_t (*Sdk_DataReportCallback)(Sdk_Channel *nChannel, uint8_t *dataBuffer, uint16_t dataSize);",
            "desc": "数据上报结果回调。数据上报完成后，SDK 通过此回调将 DK 设备的响应数据通知 MCU。",
            "params": [
                ("nChannel", "Sdk_Channel*", "通信通道"),
                ("dataBuffer", "uint8_t*", "DK 设备响应数据缓冲区"),
                ("dataSize", "uint16_t", "响应数据长度"),
            ],
            "return": "uint8_t：0x00=成功",
            "notes": "回调中 dataBuffer 指向的缓冲区在回调返回后可能被 SDK 重用，MCU 应在回调内部完成数据拷贝。"
        },
        {
            "name": "Sdk_GetLogCallback",
            "proto": "typedef void (*Sdk_GetLogCallback)(uint8_t *logBuffer, uint16_t logSize);",
            "desc": "日志输出回调。SDK 通过此回调将运行日志输出给 MCU，MCU 可将日志输出到串口或存储设备用于调试。",
            "params": [
                ("logBuffer", "uint8_t*", "日志数据缓冲区"),
                ("logSize", "uint16_t", "日志数据长度"),
            ],
            "return": "无返回值（void）。",
            "notes": "日志输出等级由宏 SDK_LOG_LEVEL 控制。设置为 SDK_LOG_LEVEL_NONE 时不会有任何日志输出。"
        },
    ]
    
    for cb in notify_callbacks:
        add_callback_detail_table(doc,
            cb["name"], cb["proto"], cb["desc"], cb["params"],
            cb["return"], cb.get("notes"))
    
    doc.add_heading("5.3  SE Agent 回调（预留）", level=2)
    
    # 接口汇总表
    doc.add_paragraph("表5-3  SE Agent 回调接口汇总")
    add_interface_summary_table(doc,
        ["回调类型（typedef）", "函数签名", "功能描述", "返回值"],
        [
            ["SEAgent_apdu", "uint8_t (*)(SEAgent_ApduParam *param)",
             "SE Agent APDU发送回调，用于TSM等场景（预留）", "0=成功"],
            ["SEAgent_ApplySeOper", "uint8_t (*)(uint16_t duration)",
             "申请SE操作权限，duration为预计操作时长ms（预留）", "0=授权, 非0=拒绝"],
            ["SEAgent_ReleaSeOper", "uint8_t (*)(void)",
             "释放SE操作权限，与ApplySeOper配对使用（预留）", "0=成功"],
            ["SEAgent_AppletUpdate", "uint8_t (*)(uint8_t result)",
             "Applet更新结果通知，result=0表示成功（预留）", "0=成功"],
        ],
        col_widths=[3.0, 5.0, 5.5, 1.5]
    )
    
    doc.add_paragraph("以下为各回调接口的详细说明：")
    
    seagent_callbacks = [
        {
            "name": "SEAgent_apdu",
            "proto": "typedef uint8_t (*SEAgent_apdu)(SEAgent_ApduParam *param);",
            "desc": "SE Agent APDU 发送回调。用于 TSM 等场景下通过 SE Agent 通道发送 APDU 指令。",
            "params": [("param", "SEAgent_ApduParam*", "APDU 参数")],
            "return": "uint8_t：0x00=成功",
            "notes": "当前为预留接口。"
        },
        {
            "name": "SEAgent_ApplySeOper",
            "proto": "typedef uint8_t (*SEAgent_ApplySeOper)(uint16_t duration);",
            "desc": "申请 SE 操作权限回调。在执行敏感操作前，SDK 通过此回调向 MCU 申请 SE 操作权限，duration 为预计操作时长（ms）。",
            "params": [("duration", "uint16_t", "预计操作时长（ms）")],
            "return": "uint8_t：0x00=授权成功，其他=拒绝",
            "notes": "当前为预留接口。MCU 可在回调中实现互斥锁机制，确保 SE 操作期间不被其他任务打断。"
        },
        {
            "name": "SEAgent_ReleaSeOper",
            "proto": "typedef uint8_t (*SEAgent_ReleaSeOper)(void);",
            "desc": "释放 SE 操作权限回调。SE 操作完成后，SDK 通过此回调通知 MCU 释放 SE 操作权限。",
            "params": [],
            "return": "uint8_t：0x00=成功",
            "notes": "当前为预留接口。与 SEAgent_ApplySeOper 配对使用。"
        },
        {
            "name": "SEAgent_AppletUpdate",
            "proto": "typedef uint8_t (*SEAgent_AppletUpdate)(uint8_t result);",
            "desc": "Applet 更新结果通知回调。Applet 更新完成后，SDK 通过此回调将更新结果通知 MCU。",
            "params": [("result", "uint8_t", "更新结果（0x00=成功）")],
            "return": "uint8_t：0x00=成功",
            "notes": "当前为预留接口。"
        },
    ]
    
    for cb in seagent_callbacks:
        add_callback_detail_table(doc,
            cb["name"], cb["proto"], cb["desc"], cb.get("params", []),
            cb["return"], cb.get("notes"))
    
    # ====================================================================
    # 第6章 典型应用流程
    # ====================================================================
    doc.add_heading("第6章  典型应用流程", level=1)
    doc.add_paragraph(
        "本章以流程图和步骤说明的方式描述 DKSDK 的典型应用场景，帮助 MCU 开发者快速理解"
        "SDK 的调用时序和集成要点。各流程中标注了 MCU 需要调用的接口以及 SDK 回调 MCU 的时机。"
    )
    
    # ---- 6.1 ----
    doc.add_heading("6.1  系统初始化流程", level=2)
    doc.add_paragraph(
        "系统上电后，MCU 需按以下步骤完成 SDK 初始化和回调注册。该流程仅需执行一次。"
    )
    add_table_with_style(doc,
        ["步骤", "操作方", "动作", "说明"],
        [
            ["1", "MCU", "实现所有必需的回调函数",
             "按第5章定义的函数签名实现 apdu_cb、send_cb、notify_auth_cb 等回调"],
            ["2", "MCU", "调用 Sdk_Register* 系列函数注册回调",
             "至少需注册 Sdk_RegisterApdu、Sdk_RegisterSend、Sdk_RegisterNotifyAuthResult"],
            ["3", "MCU", "调用 Sdk_Init()",
             "SDK 清零会话/队列，自动发送 SELECT 指令选择 SE Applet"],
            ["4", "SDK→MCU", "回调 get_version_cb（若已注册）",
             "SDK 解析 SE 响应后通过此回调通知版本信息（SDK版本/Applet版本/SEID）"],
            ["5", "MCU", "启动定时器，周期性调用 Sdk_routine()",
             "建议 1ms 周期，驱动 SDK 状态机、超时检测和响应分发"],
        ],
        col_widths=[1.0, 2.0, 5.0, 7.0]
    )
    doc.add_paragraph(
        "注意：步骤4（get_version_cb）在 Sdk_Init 流程中异步触发，由 Sdk_routine 主循环驱动完成。"
        "MCU 不应在 Sdk_Init 返回后立即假设版本信息已就绪。"
    )
    
    # ---- 6.2 ----
    doc.add_heading("6.2  数字钥匙认证流程", level=2)
    doc.add_paragraph(
        "认证是 DKSDK 的核心功能。SDK 内部通过状态机驱动 SE 与手机/卡片间的交互，MCU 仅需充当通信“桥梁”并处理结果。"
    )
    add_table_with_style(doc,
        ["步骤", "操作方", "动作", "说明"],
        [
            ["1", "MCU", "调用 Sdk_Auth(&channel)",
             "指定通信通道（NFC或BLE），发起认证。若通道忙则返回 STATUS_FAILED"],
            ["2", "SDK→MCU", "回调 apdu_cb(Select SE)",
             "SDK 要求 MCU 向 SE 发送指令以获取读卡器参数"],
            ["3", "MCU→SDK", "调用 Sdk_ApduCallback()",
             "MCU 将 SE 响应回传，SDK 内部解析随机数等参数"],
            ["4", "SDK→MCU", "回调 send_cb(Auth Req)",
             "SDK 将认证请求发送给 MCU，MCU 通过无线通道（BLE/NFC）转发给手机"],
            ["5", "MCU→SDK", "调用 Sdk_SendCallback()",
             "手机响应后，MCU 将数据回传给 SDK，SDK 进一步与 SE 交互"],
            ["6", "SDK", "多次往返校验",
             "SDK 自动完成内部认证、DK 认证和外部认证的指令闭环"],
            ["7", "SDK→MCU", "回调 notify_auth_cb",
             "流程结束：errorCode=0 表示成功，含 CardID；非 0 表示失败"],
        ],
        col_widths=[1.0, 2.0, 4.0, 8.0]
    )
    doc.add_paragraph(
        "注意：MCU 无需了解认证内部状态。只需确保在收到回调时执行硬件收发，并将结果及时回传即可。"
    )

    
    # ---- 6.3 ----
    doc.add_heading("6.3  数据上报流程", level=2)
    doc.add_paragraph(
        "MCU 可通过 Sdk_DataReport 向 DK 设备上报数据（如车辆状态、电量信息等），数据经 SE 加密后通过通信通道发送。"
    )
    add_table_with_style(doc,
        ["步骤", "操作方", "动作", "说明"],
        [
            ["1", "MCU", "调用 Sdk_DataReport(&param)",
             "填充 Sdk_DataReportParam：通道信息 + 16字节数据 + 有效长度"],
            ["2", "SDK", "通道忙检测",
             "若通道正被 RKE/Auth/GetCalibData 占用，返回 STATUS_FAILED；MCU 应在 2s 后重试"],
            ["3", "SDK", "构建 TLV + 加密",
             "SDK 将上报数据构建为 TLV 格式（Tag=0x03），经 SE 加密（Encrypt_Decrypt, type=0x01）"],
            ["4", "SDK→MCU", "回调 send_cb 发送加密数据",
             "MCU 将加密后的数据通过指定通道发送给 DK 设备"],
            ["5", "MCU→SDK", "调用 Sdk_SendCallback 回传 DK 响应",
             "DK 设备响应数据回传给 SDK"],
            ["6", "SDK→MCU", "回调 data_report_cb（若已注册）",
             "SDK 将 DK 响应数据通知 MCU，完成上报流程"],
            ["7", "SDK", "自动结束会话",
             "上报完成后 SDK 自动关闭会话；若有 pending RKE 则自动处理"],
        ],
        col_widths=[1.0, 1.8, 4.5, 7.7]
    )
    doc.add_paragraph(
        "注意事项：dataBuffer 为定长 16 字节，dataSize 指定有效数据长度（≤16）。"
        "上报数据在 SE 内部经加密处理后发送，确保数据传输安全性。"
    )
    
    # ---- 6.4 ----
    doc.add_heading("6.4  RKE 遥控钥匙流程", level=2)
    doc.add_paragraph(
        "RKE（Remote Keyless Entry）由手机端发起。SDK 负责指令解密，MCU 负责执行控制动作。"
    )
    add_table_with_style(doc,
        ["步骤", "操作方", "动作", "说明"],
        [
            ["1", "MCU→SDK", "调用 Sdk_SendCallback(RKE数据)",
             "MCU 收到手机发来的 RKE 密文消息，回传给 SDK"],
            ["2", "SDK→MCU", "回调 apdu_cb(解密请求)",
             "SDK 要求 MCU 调用 SE 硬件对密文进行解密"],
            ["3", "MCU→SDK", "调用 Sdk_ApduCallback()",
             "MCU 回传解密后的数据给 SDK"],
            ["4", "SDK→MCU", "回调 rke_cb(明文指令)",
             "SDK 解析出明文 RKEcmd，通知 MCU 执行车辆动作"],
            ["5", "MCU", "执行动作并调用 Sdk_RKECallback()",
             "MCU 执行完解锁等动作后，回传执行结果"],
            ["6", "SDK→MCU", "回调 send_cb(加密反馈)",
             "SDK 将加密后的执行结果通过 MCU 发回手机"],
        ],
        col_widths=[1.0, 2.2, 4.3, 7.5]
    )
    doc.add_paragraph(
        "并发挂起机制：若认证正在进行时收到 RKE，SDK 会自动将其挂起（Pending），并在认证结束后自动触发处理，无需 MCU 侧干预。"
    )

    
    # ---- 6.5 ----
    doc.add_heading("6.5  校准数据获取流程", level=2)
    doc.add_paragraph(
        "MCU 可通过 Sdk_GetCalibData 向 SE 请求校准数据，用于 BLE 测距等场景的校准参数获取。"
    )
    add_table_with_style(doc,
        ["步骤", "操作方", "动作", "说明"],
        [
            ["1", "MCU", "调用 Sdk_GetCalibData(&req)",
             "填充 Sdk_GetCalibReq：通道信息 + 请求数据"],
            ["2", "SDK", "通道忙检测",
             "若通道正被 RKE/Auth/DataReport 占用，返回 STATUS_FAILED"],
            ["3", "SDK", "构建 TLV + 加密发送",
             "将请求数据构建为 TLV（Tag=0x06），经 SE 加密后发送"],
            ["4", "SDK→MCU", "回调 notify_calib_cb（若已注册）",
             "校准数据就绪后通过此回调通知 MCU；MCU 需在回调内完成数据拷贝"],
        ],
        col_widths=[1.0, 1.8, 4.5, 7.7]
    )
    
    # ---- 6.6 ----
    doc.add_heading("6.6  多通道并发管理", level=2)
    doc.add_paragraph(
        "DKSDK 支持最多 SDK_MAX_SESSIONS（6）个并发会话，可同时管理多个 NFC 和 BLE 通道。"
        "每个会话拥有独立的上下文（状态机状态、超时计数器、响应缓冲区、认证参数），互不干扰。"
    )
    doc.add_paragraph("通道与会话索引的映射关系：")
    add_table_with_style(doc,
        ["通道类型", "通道ID（id[0]）", "会话索引", "说明"],
        [
            ["NFC", "0", "0", "NFC 通道 0 → 会话 0"],
            ["NFC", "1", "1", "NFC 通道 1 → 会话 1"],
            ["BLE", "0", "2", "BLE 通道 0 → 会话 2"],
            ["BLE", "1", "3", "BLE 通道 1 → 会话 3"],
            ["BLE", "2", "4", "BLE 通道 2 → 会话 4"],
            ["BLE", "3", "5", "BLE 通道 3 → 会话 5"],
        ],
        col_widths=[3.0, 3.0, 3.0, 6.0]
    )
    doc.add_paragraph(
        "并发保护机制："
    )
    doc.add_paragraph("SE 指令队列（Sdk_HwApduManager）：所有会话共享一个 SE 指令队列（深度 6），使用队列锁（queue_locked）保护并发入队/出队操作。", style='List Bullet')
    doc.add_paragraph("会话隔离：每个会话拥有独立的 apdu_resp[]、dk_apdu_resp[]、tlv_buffer[]、reader_rnd[] 等缓冲区，防止多连接数据覆盖。", style='List Bullet')
    doc.add_paragraph("通道忙检测（Sdk_IsChannelBusyFor）：同一通道上不允许 RKE/Auth/DataReport/GetCalibData 互相抢占，防止状态机混乱。", style='List Bullet')
    doc.add_paragraph("Pending RKE 机制：当通道正忙时 RKE 数据暂存，当前操作完成后自动执行，保证指令不丢失。", style='List Bullet')
    
    # ---- 6.7 ----
    doc.add_heading("6.7  超时与异常处理", level=2)
    doc.add_paragraph(
        "DKSDK 内置多层超时保护机制，由 Sdk_routine 主循环统一检测和处理。"
    )
    add_table_with_style(doc,
        ["超时类型", "检测阈值", "触发条件", "处理方式"],
        [
            ["会话总超时", "SESSION_TIMEOUT_TICKS（3000）",
             "会话 overall_ticks 超过阈值（默认 3s @ 1ms 周期）",
             "关闭会话，通过 notify_auth_cb 报告 ERR_AUTH_TIMEOUT（0x18）"],
            ["SE 单指令超时", "SE_CMD_TIMEOUT_TICKS（1000）",
             "SE 指令队列头部指令超过 1s 无响应",
             "强制出队，关闭会话，报告 ERR_COMMAND_TIMEOUT（0x19）"],
            ["等待对端超时", "PEER_STEP_TIMEOUT_TICKS（1000）",
             "会话 is_waiting_peer=1 且超过 1s 未收到对端回调",
             "关闭会话，报告 ERR_COMMAND_TIMEOUT（0x19）"],
            ["APDU 响应溢出", "—",
             "Sdk_ApduCallback 被调用时上一帧数据尚未被主循环消费",
             "返回 STATUS_FAILED，数据丢弃；MCU 应检查调用频率"],
            ["队列锁超时", "重试 100 次",
             "入队/出队时超过 100 次未获取到 queue_locked 锁",
             "放弃本次操作，输出错误日志"],
        ],
        col_widths=[2.5, 3.5, 4.5, 4.5]
    )
    doc.add_paragraph(
        "MCU 侧建议：若 MCU 改变 Sdk_routine 调用周期，需同步修改 SESSION_TIMEOUT_TICKS、"
        "SE_CMD_TIMEOUT_TICKS、PEER_STEP_TIMEOUT_TICKS 三个宏（位于 DKSdk.c 顶部），"
        "确保实际超时时间符合设计预期。"
    )
    
    # ---- 6.8 错误码速查 ----
    doc.add_heading("6.8  错误码速查", level=2)
    doc.add_paragraph(
        "以下汇总 SDK 所有错误码及其含义和 MCU 建议处理方式，便于快速定位问题。"
    )
    add_table_with_style(doc,
        ["错误码", "常量名", "含义", "MCU建议处理"],
        [
            ["0x00", "（成功）", "操作成功完成", "正常继续"],
            ["0x10", "ERR_AUTH_BASE", "认证错误基准值（内部使用）", "—"],
            ["0x11", "ERR_AUTH_SELECT_SE_FAIL", "SE Applet 选择失败", "检查 SE 硬件连接，重试或重新初始化"],
            ["0x12", "ERR_AUTH_SELECT_DK_FAIL", "DK Applet 选择失败", "检查 DK 设备是否在通信范围内"],
            ["0x13", "ERR_AUTH_GPD_FAIL", "GPD 指令失败", "检查 DK 设备兼容性"],
            ["0x14", "ERR_AUTH_INTERNAL_FAIL", "SE 内部认证失败", "检查 SE 密钥配置"],
            ["0x15", "ERR_AUTH_DK_AUTH_FAIL", "DK 认证失败", "DK 设备可能未授权，提示用户"],
            ["0x16", "ERR_AUTH_EXTERNAL_FAIL", "外部认证失败", "DK 与 SE 密钥不匹配"],
            ["0x17", "ERR_AUTH_NO_PERMISSION", "无权限", "DK 设备未被授权访问该车辆"],
            ["0x18", "ERR_AUTH_TIMEOUT", "认证总超时（3s）", "检查通信链路质量，提示用户靠近设备"],
            ["0x19", "ERR_COMMAND_TIMEOUT", "指令超时（1s）", "SE 或 DK 响应超时，重试"],
            ["0x1A", "ERR_NOT_INIT", "SDK 未初始化", "调用 Sdk_Init 后再操作"],
            ["0x21", "ERR_PARSE_CRC_FAIL", "BLE 帧 CRC 校验失败", "检查 BLE 链路质量，可能存在干扰"],
            ["0x22", "ERR_PARSE_FRAME_TOO_SHORT", "BLE 帧长度不足", "检查 ICCE 协议栈版本兼容性"],
            ["0x23", "ERR_PARSE_SOF_WRONG", "BLE 帧起始符错误", "帧同步丢失，检查数据流完整性"],
            ["0x24", "ERR_PARSE_MESSAGEID_WRONG", "消息 ID 不合法", "协议版本不匹配"],
            ["0x25", "ERR_PARSE_COMMANDID_WRONG", "命令 ID 不合法", "协议版本不匹配"],
            ["0x26", "ERR_PARSE_TLV_WRONG", "TLV 解析失败", "数据格式异常"],
            ["0x27", "ERR_PARSE_SW_WRONG", "状态字 SW 错误", "SE 返回异常状态"],
        ],
        col_widths=[1.0, 3.5, 4.5, 6.0]
    )
    
    # ---- 6.9 日志调试 ----
    doc.add_heading("6.9  日志调试指南", level=2)
    doc.add_paragraph(
        "DKSDK 内置分级日志系统，通过宏 SDK_LOG_LEVEL 控制输出等级。日志通过 get_log_cb 回调输出给 MCU。"
    )
    add_table_with_style(doc,
        ["日志等级", "宏值", "输出内容", "适用场景"],
        [
            ["NONE", "SDK_LOG_LEVEL_NONE（0）", "无任何日志输出", "量产固件"],
            ["ERROR", "SDK_LOG_LEVEL_ERROR（1）", "错误信息 + 超时通知", "日常调试（默认）"],
            ["DEBUG", "SDK_LOG_LEVEL_DEBUG（2）", "所有日志：状态切换、数据Hex Dump、帧解析详情", "开发调试"],
        ],
        col_widths=[2.0, 4.0, 5.0, 4.0]
    )
    doc.add_paragraph(
        "开发调试建议：开发阶段将 SDK_LOG_LEVEL 设为 SDK_LOG_LEVEL_DEBUG，注册 get_log_cb 将日志输出到串口。"
        "量产阶段改为 SDK_LOG_LEVEL_NONE 或 SDK_LOG_LEVEL_ERROR 以减少性能开销。"
    )
    
    # ---- 6.10 ----
    doc.add_heading("6.10 重要注意事项", level=2)
    doc.add_paragraph("在集成 DKSDK 时，MCU 开发者请务必注意以下事项：")
    doc.add_paragraph("（1）回调非阻塞：apdu_cb 和 send_cb 必须立即返回，不可有耗时循环。", style='List Bullet')
    doc.add_paragraph("（2）调度频率：必须以 1ms 周期稳定调用 Sdk_routine()，否则会影响超时检测。", style='List Bullet')
    doc.add_paragraph("（3）数据拷贝：回调内提供的 dataBuffer 仅在回调期间有效，如需后续使用请进行拷贝。", style='List Bullet')
    doc.add_paragraph("（4）并发限制：同一通道上，Auth、DataReport 和 GetCalibData 互斥，忙时会返回失败。", style='List Bullet')
    doc.add_paragraph("（5）连接释放：蓝牙或 NFC 断开时，必须调用 Sdk_Release 释放 SDK 内部会话。", style='List Bullet')
    doc.add_paragraph("（6）通道 ID：请确保为不同的物理连接（如 4 路蓝牙）分配唯一的 channel.id。", style='List Bullet')

    
    # ====================================================================
    # 附录：使用示例
    # ====================================================================
    doc.add_page_break()
    doc.add_heading("附录A  MCU 集成示例代码", level=1)
    doc.add_paragraph("以下为 MCU 集成 DKSDK 的典型初始化和使用示例代码：")
    
    example_code = """// =====================================================
// MCU 集成 DKSDK 示例代码
// =====================================================

#include "DKSdk.h"

// ---- 1. 实现 MCU→SDK 回调函数 ----

// APDU 发送回调：将指令通过 SE 硬件发送
uint8_t My_ApduSend(Sdk_ApduParam *param) {
    // 通过 SE 硬件接口发送 APDU 指令
    SE_HW_SendApdu(param->dataBuffer, param->dataSize);
    return 0;
}

// 数据发送回调：将数据通过指定通道发送
uint8_t My_DataSend(Sdk_SendParam *param) {
    if (param->channel.channel_type == SDK_CHANNEL_TYPE_BLE) {
        BLE_SendData(param->channel.id[0], param->dataBuffer, param->dataSize);
    } else {
        NFC_SendData(param->dataBuffer, param->dataSize);
    }
    return 0;
}

// 认证结果通知回调
uint8_t My_AuthResult(Sdk_NotifyAuthResultParam *param) {
    if (param->errorCode == 0) {
        printf("Auth SUCCESS! CardID: ...");
    } else {
        printf("Auth FAILED! Error: 0x%02X", param->errorCode);
    }
    return 0;
}

// RKE 指令回调
uint8_t My_RKE(Sdk_RKEParam *param) {
    // 执行车辆控制：解锁/锁车
    Vehicle_ExecuteRKE(param->RKEcmd);
    return 0;
}

// 版本信息回调
uint8_t My_Version(Sdk_VersionParam *param) {
    printf("SDK Ver: %02X.%02X.%02X.%02X",
           param->SDK_Version[0], param->SDK_Version[1],
           param->SDK_Version[2], param->SDK_Version[3]);
    return 0;
}

// ---- 2. 初始化 SDK ----
void MCU_Init(void) {
    // 注册所有回调
    Sdk_RegisterApdu(My_ApduSend);
    Sdk_RegisterSend(My_DataSend);
    Sdk_RegisterNotifyAuthResult(My_AuthResult);
    Sdk_RegisterRKE(My_RKE);
    Sdk_RegisterGetVersionCallback(My_Version);
    
    // 初始化 SDK（会自动发送 SE Applet 选择指令）
    Sdk_Init();
}

// ---- 3. 主循环（1ms 周期调用） ----
void MCU_MainLoop(void) {
    Sdk_routine();  // 驱动 SDK 状态机
}

// ---- 4. SE APDU 响应中断处理 ----
void SE_ISR_Handler(uint8_t result, uint8_t *resp, uint16_t len) {
    Sdk_ApduCallback(result, resp, len);
}

// ---- 5. BLE 数据接收回调 ----
void BLE_RecvCallback(uint8_t channelId, uint8_t *data, uint16_t len) {
    Sdk_Channel ch = {SDK_CHANNEL_TYPE_BLE, 0x01, {channelId,0}, 1};
    Sdk_SendCallback(&ch, 0, data, len);
}

// ---- 6. 发起认证 ----
void StartAuth(uint8_t bleChanId) {
    Sdk_Channel ch = {SDK_CHANNEL_TYPE_BLE, 0x01, {bleChanId,0}, 1};
    Sdk_Auth(&ch);
}"""
    
    add_code_block(doc, example_code, font_size=7.5)
    
    # ====================================================================
    # 保存文档
    # ====================================================================
    output_path = os.path.join(os.path.dirname(__file__), "DKSDK接口文档_Final.docx")
    try:
        doc.save(output_path)
        print(f"成功生成最终文档：{output_path}")
    except Exception as e:
        print(f"生成失败：{str(e)}")
    return output_path



if __name__ == "__main__":
    generate_document()
